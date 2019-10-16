#!/usr/bin/env python3
# -*-coding:utf-8-*-
import os
from configparser import ConfigParser
from io import StringIO
from io import open
from concurrent.futures import ProcessPoolExecutor
from pdfminer.pdfinterp import process_pdf
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from docx import Document
import re
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter

def pdf2Txt(file_path,outfile=None):
    with open(file_path, 'rb') as file:
        resource_manager = PDFResourceManager()
        return_str = StringIO()
        lap_params = LAParams()

        device = TextConverter(
            resource_manager, return_str, laparams=lap_params)
        process_pdf(resource_manager, device, file)
        device.close()

        content = return_str.getvalue()
        content = re.sub('\n{2,}','\n',content)
        if outfile is not None:
            with open(outfile,'w') as f:
                f.write(content)
        return_str.close()
        return content


def save_text_to_word(content, file_path):
    doc = Document()
    for line in content.split('\n'):
        paragraph = doc.add_paragraph()
        paragraph.add_run(remove_control_characters(line))
    doc.save(file_path)


def remove_control_characters(content):
    mpa = dict.fromkeys(range(32))
    return content.translate(mpa)


def pdf2Word(pdf_file_path, word_file_path):
    content = pdf2Txt(pdf_file_path)
    save_text_to_word(content, word_file_path)


def main():
    config_parser = ConfigParser()
    config_parser.read('config.cfg')
    config = config_parser['default']

    tasks = []
    with ProcessPoolExecutor(max_workers=int(config['max_worker'])) as executor:
        for f in os.listdir(config['pdf_folder']):
            extension_name = os.path.splitext(f)[1]
            if extension_name != '.pdf':
                continue
            file_name = os.path.splitext(f)[0]
            pdf_file = config['pdf_folder'] + '/' + f
            word_file = config['word_folder'] + '/' + file_name + '.docx'
            print('正在处理: ', f)
            result = executor.submit(pdf2Word, pdf_file, word_file)
            tasks.append(result)
    while True:
        exit_flag = True
        for task in tasks:
            if not task.done():
                exit_flag = False
        if exit_flag:
            print('完成')
            exit(0)



#一个文件夹下的所有pdf文档转换成txt
def fileTotxt(fileDir):
    files=os.listdir(fileDir)
    tarDir=fileDir+'txt'
    if not os.path.exists(tarDir):
        os.mkdir(tarDir)
    replace=re.compile(r'\.pdf',re.I)
    for f in files:
        filePath=fileDir+'\\'+f
        outPath=tarDir+'\\'+replace.sub('',f)+'.txt'
        pdf2Txt(filePath,outPath)
        print("Saved "+outPath)

if __name__ == '__main__':
    pdf2Txt('SignalFile.pdf', 'test.txt')
    #fileTotxt('这里是目录的路径')

